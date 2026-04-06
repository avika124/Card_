"""
docx_generator.py — Generate color-coded Word documents from compliance findings.

Color coding:
  Yellow  (#FFF2CC) — Corrected / Changed provisions
  Red     (#FFCCCC) — Deleted / removed text
  Green   (#CCFFCC) — Added / new provisions
  Gray    (#F2F2F2) — Unchanged
"""

import io
from typing import Optional
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Color constants ──────────────────────────────────────────────────────────
NAVY    = RGBColor(0x1C, 0x2D, 0x4F)
GRAY    = RGBColor(0x5A, 0x5A, 0x5A)
BLACK   = RGBColor(0x00, 0x00, 0x00)
RED     = RGBColor(0xC0, 0x00, 0x00)
GREEN   = RGBColor(0x00, 0x70, 0x00)
AMBER   = RGBColor(0xBF, 0x8F, 0x00)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)

BG_UNCHANGED = "F2F2F2"
BG_CHANGED   = "FFF2CC"
BG_DELETED   = "FFCCCC"
BG_ADDED     = "CCFFCC"
BG_NAVY      = "1C2D4F"
BG_SUMMARY   = "E8E8E8"

SEVERITY_BG = {
    "high":   "FFCCCC",
    "medium": "FFF2CC",
    "low":    "E8F5E9",
    "pass":   "CCFFCC",
}
SEVERITY_COLOR = {
    "high":   RED,
    "medium": AMBER,
    "low":    RGBColor(0x2E, 0x7D, 0x32),
    "pass":   GREEN,
}


# ── Helpers ──────────────────────────────────────────────────────────────────

def _shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove existing shd
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell, color: str = "CCCCCC"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.insert(0, tcBorders)  # insert before shd


def _set_font(run, size: int = 10, bold: bool = False, italic: bool = False,
              color: Optional[RGBColor] = None, strike: bool = False):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    if strike:
        run.font.strike = True


def _para(doc, space_before: int = 0, space_after: int = 6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def _add_text(doc, text: str, size: int = 10, bold: bool = False,
              italic: bool = False, color: Optional[RGBColor] = None,
              space_before: int = 0, space_after: int = 6):
    p = _para(doc, space_before, space_after)
    r = p.add_run(text)
    _set_font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def _heading(doc, text: str, level: int = 1):
    sizes = {1: 14, 2: 12, 3: 10}
    return _add_text(doc, text, size=sizes.get(level, 10), bold=True,
                     color=NAVY, space_before=10, space_after=5)


def _table_header_row(table, headers: list[str]):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        _shade_cell(cell, BG_NAVY)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(h)
        _set_font(r, bold=True, color=WHITE, size=9)


def _finding_row(table, num: str, pri: str, regulation: str,
                 severity: str, issue: str, detail: str,
                 excerpt: str, recommendation: str):
    row = table.add_row()
    cells = row.cells
    bg = SEVERITY_BG.get(severity, BG_UNCHANGED)
    sc = SEVERITY_COLOR.get(severity, BLACK)

    for c in cells:
        _shade_cell(c, bg)
        _set_cell_borders(c, "AAAAAA")

    # Col 0: number
    p = cells[0].paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    _set_font(p.add_run(num), size=8, bold=True)

    # Col 1: severity pill
    p = cells[1].paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    _set_font(p.add_run(severity.upper()), size=8, bold=True, color=sc)

    # Col 2: regulation
    p = cells[2].paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    _set_font(p.add_run(regulation), size=8)

    # Col 3: issue + detail + excerpt + recommendation
    p = cells[3].paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    _set_font(p.add_run(issue), size=8, bold=True)

    if detail:
        p2 = cells[3].add_paragraph()
        p2.paragraph_format.space_before = Pt(3)
        p2.paragraph_format.space_after = Pt(2)
        _set_font(p2.add_run(detail), size=8, color=GRAY)

    if excerpt:
        p3 = cells[3].add_paragraph()
        p3.paragraph_format.space_before = Pt(3)
        p3.paragraph_format.space_after = Pt(2)
        r = p3.add_run("Excerpt: ")
        _set_font(r, size=8, bold=True, italic=True)
        r2 = p3.add_run(f'"{excerpt}"')
        _set_font(r2, size=8, italic=True, color=GRAY)

    if recommendation:
        p4 = cells[3].add_paragraph()
        p4.paragraph_format.space_before = Pt(4)
        p4.paragraph_format.space_after = Pt(2)
        r = p4.add_run("→ Recommendation: ")
        _set_font(r, size=8, bold=True, color=GREEN)
        r2 = p4.add_run(recommendation)
        _set_font(r2, size=8, color=GREEN)


# ── Main generator ────────────────────────────────────────────────────────────

def generate_compliance_docx(
    findings_result: dict,
    document_name: str = "Analyzed Document",
    original_text_excerpt: str = "",
) -> bytes:
    """
    Generate a color-coded compliance report Word document.

    Args:
        findings_result: dict from compliance.check_text / check_image / check_file
        document_name: display name of the analyzed document
        original_text_excerpt: first ~500 chars of original text for context

    Returns:
        bytes of the .docx file
    """
    doc = Document()

    # Page setup — US Letter, 1" margins
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(section, attr, Inches(1))

    # ── Cover ─────────────────────────────────────────────────────────────────
    _add_text(doc, "CREDIT CARD COMPLIANCE CHECKER", size=11, bold=True, color=NAVY, space_after=2)
    _add_text(doc, "AI-Powered Regulatory Analysis — Anthropic Claude", size=9, color=GRAY, space_after=4)
    _add_text(doc, f"Document analyzed: {document_name}", size=10, italic=True, color=GRAY, space_after=2)

    overall = findings_result.get("overall_risk", "unknown").upper()
    overall_color = SEVERITY_COLOR.get(findings_result.get("overall_risk", "low"), BLACK)
    p = _para(doc, space_after=8)
    r = p.add_run(f"Overall Risk: {overall}")
    _set_font(r, size=13, bold=True, color=overall_color)

    # ── Color legend ──────────────────────────────────────────────────────────
    _heading(doc, "Color-Coding Legend", level=2)
    legend = doc.add_table(rows=4, cols=2)
    legend.style = "Table Grid"
    legend.columns[0].width = Inches(1.5)
    legend.columns[1].width = Inches(6.0)

    legend_data = [
        (BG_UNCHANGED, "UNCHANGED",         "Original content — no issue found"),
        (BG_CHANGED,   "CORRECTED/CHANGED", "Provision was wrong or inconsistent — correction shown"),
        (BG_DELETED,   "HIGH RISK / DELETED","Significant compliance gap — removal or replacement required"),
        (BG_ADDED,     "ADDED / REQUIRED",  "Missing provision — must be added to achieve compliance"),
    ]
    for i, (bg, label, desc) in enumerate(legend_data):
        _shade_cell(legend.rows[i].cells[0], bg)
        _shade_cell(legend.rows[i].cells[1], bg)
        for c in legend.rows[i].cells:
            _set_cell_borders(c)
        p0 = legend.rows[i].cells[0].paragraphs[0]
        p0.paragraph_format.space_before = Pt(3)
        p0.paragraph_format.space_after  = Pt(3)
        _set_font(p0.add_run(label), size=9, bold=True)
        p1 = legend.rows[i].cells[1].paragraphs[0]
        p1.paragraph_format.space_before = Pt(3)
        p1.paragraph_format.space_after  = Pt(3)
        _set_font(p1.add_run(desc), size=9)

    doc.add_paragraph()

    # ── Executive summary ─────────────────────────────────────────────────────
    _heading(doc, "Executive Summary", level=1)
    _add_text(doc, findings_result.get("summary", ""), size=10, space_after=8)

    # ── Stats row ─────────────────────────────────────────────────────────────
    findings = findings_result.get("findings", [])
    counts = {"high": 0, "medium": 0, "low": 0, "pass": 0}
    for f in findings:
        sev = f.get("severity", "low")
        counts[sev] = counts.get(sev, 0) + 1

    stat_table = doc.add_table(rows=2, cols=4)
    stat_table.style = "Table Grid"
    stat_headers = ["HIGH RISK", "MEDIUM RISK", "LOW RISK", "PASSING"]
    stat_bgs     = [BG_DELETED, BG_CHANGED, "E8F5E9", BG_ADDED]
    stat_vals    = [counts["high"], counts["medium"], counts["low"], counts["pass"]]
    stat_colors  = [RED, AMBER, RGBColor(0x2E,0x7D,0x32), GREEN]

    for i in range(4):
        _shade_cell(stat_table.rows[0].cells[i], stat_bgs[i])
        _shade_cell(stat_table.rows[1].cells[i], stat_bgs[i])
        for r in range(2):
            _set_cell_borders(stat_table.rows[r].cells[i])

        p0 = stat_table.rows[0].cells[i].paragraphs[0]
        p0.paragraph_format.space_before = Pt(4)
        p0.paragraph_format.space_after  = Pt(2)
        p0.alignment = 1  # center
        _set_font(p0.add_run(str(stat_vals[i])), size=20, bold=True, color=stat_colors[i])

        p1 = stat_table.rows[1].cells[i].paragraphs[0]
        p1.paragraph_format.space_before = Pt(2)
        p1.paragraph_format.space_after  = Pt(4)
        p1.alignment = 1
        _set_font(p1.add_run(stat_headers[i]), size=8, bold=True, color=GRAY)

    doc.add_paragraph()

    # ── Findings table ────────────────────────────────────────────────────────
    _heading(doc, "Regulatory Findings", level=1)

    findings_table = doc.add_table(rows=1, cols=4)
    findings_table.style = "Table Grid"
    findings_table.columns[0].width = Inches(0.4)
    findings_table.columns[1].width = Inches(0.7)
    findings_table.columns[2].width = Inches(1.4)
    findings_table.columns[3].width = Inches(5.3)
    _table_header_row(findings_table, ["#", "Severity", "Regulation", "Finding & Recommendation"])

    # Sort: high → medium → low → pass
    order = {"high": 0, "medium": 1, "low": 2, "pass": 3}
    sorted_findings = sorted(findings, key=lambda f: order.get(f.get("severity", "low"), 2))

    for i, f in enumerate(sorted_findings, 1):
        _finding_row(
            findings_table,
            num=str(i),
            pri="",
            regulation=f.get("regulation", ""),
            severity=f.get("severity", "low"),
            issue=f.get("issue", ""),
            detail=f.get("detail", ""),
            excerpt=f.get("excerpt", ""),
            recommendation=f.get("recommendation", ""),
        )

    doc.add_paragraph()

    # ── Original content excerpt ───────────────────────────────────────────────
    if original_text_excerpt:
        _heading(doc, "Analyzed Content (Excerpt)", level=2)
        p = _para(doc, space_after=4)
        r = p.add_run(original_text_excerpt[:800] + ("..." if len(original_text_excerpt) > 800 else ""))
        _set_font(r, size=8, italic=True, color=GRAY)
        doc.add_paragraph()

    # ── Footer note ───────────────────────────────────────────────────────────
    p = _para(doc, space_before=12, space_after=4)
    r = p.add_run("This report was generated by the Credit Card Compliance Checker Agent using Anthropic Claude. "
                  "It does not constitute legal advice. Review with qualified legal and compliance counsel before "
                  "taking action on any finding.")
    _set_font(r, size=8, italic=True, color=GRAY)

    # ── Serialize to bytes ────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
